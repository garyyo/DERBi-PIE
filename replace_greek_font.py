from collections import defaultdict
import unicodedata as ud
from docx import Document
from docx.document import Document as DocType
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import pyperclip
import difflib
import regex as re
from tqdm import tqdm


# This is incomplete, others may need to be found and manually identified.
greek_font2unicode = {
    'a': 'α', 'b': 'β', 'c': 'χ', 'd': 'δ', 'e': 'ε', 'f': 'φ', 'g': 'γ', 'h': 'η', 'i': 'ι', 'j': 'ς', 'k': 'κ', 'l': 'λ', 'm': 'μ', 'n': 'ν', 'o': 'ο',
    'p': 'π', 'q': 'θ', 'r': 'ρ', 's': 'σ', 't': 'τ', 'u': 'υ', 'v': 'ᾳ', 'w': 'ω', 'x': 'ξ', 'y': 'ψ', 'z': 'ζ', 'A': 'Α', 'B': 'Β', 'C': 'Χ', 'D': 'Δ',
    'E': 'Ε', 'F': 'Φ', 'G': 'Γ', 'H': 'Η', 'I': 'Ι', 'J': 'ῳ', 'K': 'Κ', 'L': 'Λ', 'M': 'Μ', 'N': 'Ν', 'O': 'Ο', 'P': 'Π', 'Q': 'Θ', 'R': 'Ρ', 'S': 'Σ',
    'T': 'Τ', 'U': 'Υ', 'V': 'ῃ', 'W': 'Ω', 'X': 'Ξ', 'Y': 'Ψ', 'Z': 'Ζ'
}
# the command I use to turn 2 copied columns in excel into the dictionary update value below.
#
greek_font2unicode.update({
    '#': 'ϝ', "'": ' ̓', '(': '(', ')': ')', '*': '*', ',': ',', '-': '-', '.': '.', ':': '·', '?': '?', '[': '[', ']': ']', '`': ' ̔', '\x8e': 'ϊ', '¢': 'ἀ',
    '£': 'ά', '¥': 'ἄ', '«': 'ἆ', '¹': 'ἡ', 'º': 'ἠ', '»': 'ή', 'À': 'ἢ', 'Ð': 'ὁ', 'Ñ': 'ὀ', 'Ù': 'ὐ', 'Ý': 'ὺ', 'à': 'ῦ', 'á': 'ὗ', 'â': 'ὖ', 'ã': 'ϋ',
    'ç': 'ὠ', 'ê': 'ὤ', 'ð': 'ὦ', 'õ': 'ᾤ', 'ù': 'ῷ', 'š': 'έ', 'ˆ': 'ὶ', '˜': 'ἑ', '’': '῏', '“': '῞', '”': '῎', '„': 'ἰ', '†': 'ἵ', '‡': 'ἴ', '‹': 'ῖ', '›': 'ἕ',
    '™': 'ἐ', '∙': 'ῥ', ' ': ' ', '\t': '\t',
})

# you cant directly write a lot of these characters in the editor, so instead their encoded versions are saved here and decoded at runtime.
titus_font2unicode = {b'\xee\x96\xb4': b'm\xcc\x81\xcc\xa5', b'\xee\x93\x9d': b'e\xcc\xaf', b'\xee\x90\x8a': b'\xc4\x81\xcc\x81', b'\xee\xb7\xb1': b'\xc9\x99', b'\xee\x95\x88': b'i\xcc\xaf', b'\xee\x93\xbb': b'g\xcc\x82', b'\xee\x90\x89': b'\xc4\x81\xcc\x80', b'\xee\x97\xae': b'n\xcc\xa5', b'\xee\x9c\xa7': b'u\xcc\xaf', b'\xee\x95\x82': b'i\xcc\x8a', b'\xee\x9a\xa3': b'r\xcc\xa5', b'\xee\x90\x93': b'\xc4\x81\xcc\x8a', b'\xee\x94\xa1': b'h\xcc\x81', b'\xee\x95\xb7': b'k\xca\xb7', b'\xee\x95\xaa': b'k\xcc\x91', b'\xee\x96\xbe': b'm\xcc\x86\xcc\x87', b'\xee\x9a\x90': b'r\xcc\x81\xcc\xa3', b'\xee\x96\x80': b'k\xca\xbf', b'\xee\xac\x80': b'\xce\xb1\xcc\x84\xcc\x81', b'\xee\x92\xb7': b'e\xcc\x84\xcc\x86', b'\xee\x94\xb5': b'\xc4\xab\xcc\x81', b'\xee\x9b\x81': b'\xc5\xa1\xcc\xa3', b'\xee\x91\xb1': b'c\xca\xbf', b'\xee\x9a\x94': b'r\xcc\x83', b'\xee\x9c\x8b': b'u\xcc\x84\xcc\x86', b'\xee\x95\xbb': b'k\xca\xbf', b'\xee\x91\xae': b'c\xca\xbf', b'\xee\x93\x89': b'e\xcc\x87\xcc\x83', b'\xee\x9c\x89': b'u\xcc\x84\xcc\x81', b'\xee\x9b\xb3': b't\xcc\xb0', b'\xee\x90\x90': b'a\xcc\x84\xcc\x86', b'\xee\xae\x81': b'\xe1\xbf\xa0\xcc\x81', b'\xee\x9c\x8a': b'u\xcc\x84\xcc\x83', b'\xee\x9c\x85': b'u\xcc\xa8\xcc\x83', b'\xee\x94\xb7': b'i\xcc\x84\xcc\x86', b'\xee\x90\x88': b'a\xcc\xa8\xcc\x83', b'\xee\x97\x85': b'm\xcc\xa5', b'\xee\x96\xa4': b'l\xcc\xa5', b'\xee\x98\x9b': b'o\xcc\x84\xcc\x86', b'\xee\x94\x8b': b'g\xca\xb7', b'\xee\x93\x8a': b'e\xcc\x87\xcc\x81', b'\xee\x96\xb7': b'm\xcc\x83', b'\xee\x93\xa6': b'\xc4\x99', b'\xee\x92\xbf': b'\xc4\x99\xcc\x84', b'\xee\x8a\xa3': b'R\xcc\xaa', b'\xee\x98\x97': b'o\xcc\x82', b'\xee\x9f\xa9': b'\xc9\x99\xcc\x81', b'\xee\x9f\xac': b'\xc9\x99\xcc\x84', b'\xee\x9b\xb9': b't\xca\xbf', b'\xee\x9a\x8d': b'r\xcc\x80', b'\xee\x9a\x92': b'r\xcc\x91', b'\xee\x92\xbc': b'\xc4\x99\xcc\x84', b'\xee\x92\xa1': b'\xc4\x99\xcc\x82', b'\xee\x9d\xa9': b'x\xe1\xb5\x9b', b'\xee\x9e\xa4': b'\xc4\xb1\xcc\x87\xcc\x83', b'\xee\x96\x82': b'k\xcd\xa1\xe1\xb5\xa4', b'\xee\x9e\xbd': b'\xc5\x8b\xe1\xb5\x9b', b'\xee\x9d\x9c': b'x\xcc\x81', b'\xee\x96\x94': b'l\xcc\x83', b'\xee\x9b\xb6': b't\xca\xbf', b'\xee\x90\x98': b'a\xcc\xa8\xcc\x87', b'\xee\x91\xa0': b'\xc4\x8d\xca\xbf', b'\xee\x98\xb8': b'o\xcc\x82\xcc\xa3', b'\xee\x93\x97': b'e\xcc\xa8\xcc\x82', b'\xee\x96\x92': b'l\xcc\x82', b'\xee\x91\xaf': b'c\xe2\x80\x99', b'\xee\x9a\x91': b'r\xcc\x81\xcc\xa5', b'\xee\x92\xa0': b'e\xcc\x82\xcc\xa7', b'\xee\xac\x8a': b'\xe1\xbe\xb0\xcc\x81', b'\xee\x96\x91': b'l\xcc\x81\xcc\xa5', b'\xee\xad\x80': b'\xe1\xbf\x91\xcc\x81', b'\xee\x92\x99': b'\xc4\x99\xcc\x81', b'\xee\x93\x8b': b'e\xcc\x88\xcc\x81', b'\xee\x9a\x96': b'r\xcc\x87', b'\xee\x92\xbe': b'e\xcc\x84\xcc\xa7'}
# we manually found a couple more later. Ideally I would normalize these
# anton: normalize these maybe eventually if it causes issues?
titus_manually_found = {b'\xee\x90\x84': b'\xc4\x85\xcc\x81', b'\xee\x90\x8e': b'\xc4\x81\xcc\x83', b'\xee\x90\xb8': b'\xc3\xa6\xcc\x83', b'\xee\x91\x9d': b'\xc4\x8d\xe2\x80\x98', b'\xee\x92\xa7': b'\xc4\x99\xcc\x83', b'\xee\x92\xac': b'\xc4\x99\xcc\x84\xcc\x80', b'\xee\x92\xb3': b'e\xcc\x84\xcc\x83', b'\xee\x94\xaf': b'\xc4\xaf\xcc\x81', b'\xee\x95\x93': b'j\xcc\x81', b'\xee\x95\xaf': b'k\xcc\xb9', b'\xee\x97\x95': b'n\xcc\xa5\xcc\x81', b'\xee\x99\xb2': b'p\xe2\x80\x98', b'\xee\x9a\x91': b'r\xcc\xa5\xcc\x81', b'\xee\x9a\x97': b'r\xcc\xa3\xcc\x84\xcc\x81', b'\xee\x9e\xb3': b'\xc5\x8b\xcc\x81', b'\xee\x9f\xae': b'\xc9\x99\xcc\x84\xcc\x86', b'\xee\xac\x84': b'\xce\xb1\xcc\x84\xcc\x93', b'\xee\xac\x85': b'\xce\xb1\xcc\x84\xcc\x94', b'\xee\xac\x87': b'\xe1\xbc\x84\xcc\x85', b'\xee\xac\xa4': b'\xce\xb5\xcc\x84', b'\xee\xac\xa6': b'\xce\xb5\xcc\x84\xcc\x81', b'\xee\xac\xb9': b'\xce\xb9\xcc\x84\xcc\x81', b'\xee\xad\xba': b'\xcf\x85\xcc\x84\xcc\x81'}
titus_font2unicode.update(titus_manually_found)

# we are no longer using the decoded version because it just leads to weirdness as we need char by char level of counting.
# titus_font2unicode = {font.decode(): unicode.decode() for font, unicode in titus_font2unicode.items()}

# compile the regex first for efficiency.
titus_pattern = re.compile(b'|'.join(re.escape(key) for key in titus_font2unicode.keys()))


def replace_using_regex_titus(text, replacements):
    # Encode the text to bytes
    byte_text = text.encode('utf-8')

    # Create a regex pattern that matches any of the keys
    # anton: this is compiled earlier to avoid needing to do it a ton of times.

    # Function to perform the replacement
    def replace_match(match):
        return replacements[match.group(0)]

    # Perform the replacement on bytes
    replaced_byte_text = titus_pattern.sub(replace_match, byte_text)

    # Decode the bytes back to string
    return replaced_byte_text.decode('utf-8')


def find_missing_greek(doc):
    # only get the segments that are in greek
    greek_segments_list = [[run for run in p.runs if run.font.name == "Greek"] for p in doc.paragraphs]
    # remove empty lists and get the text out
    greek_segments_list = ["".join([segment.text for segment in run]) for run in greek_segments_list if len(run) != 0]

    # identify the characters used in the Greek font in the doc and the Greek font ones that we already know.
    greek_chars = set("".join(greek_segments_list))
    existing_chars = set(greek_font2unicode.keys())

    missing_chars = greek_chars - existing_chars - {"\t", " "}
    paste_able = "\n".join(sorted(missing_chars))
    # pyperclip.copy(paste_able)
    breakpoint()
    pass


def iter_block_items(parent):
    # https://github.com/python-openxml/python-docx/issues/40
    """
    Yield each paragraph and table child within *parent*, in document order.
    Each returned value is an instance of either Table or Paragraph. *parent*
    would most commonly be a reference to a main Document object, but
    also works for a _Cell object, which itself can contain paragraphs and tables.
    """
    if isinstance(parent, DocType):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    # print('parent_elm: '+str(type(parent_elm)))
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            # yield Table(child, parent)  # No recursion, return tables as tables
            table = Table(child, parent)  # Use recursion to return tables as paragraphs
            for row in table.rows:
                for cell in row.cells:
                    yield from iter_block_items(cell)


def find_bad_titus(doc1, doc2):
    broken2fixed = {
        # '\ue727': b'u\xcc\xaf'.decode(),
        # '\ueb81': b'\xcf\x85\xcc\x86\xcc\x81'.decode(),
        # '\ue548': b'\xcc\xaf'.decode(),
        '\ue5b4': b'm\xcc\x81\xcc\xa5'.decode(),
        '\ue4dd': b'e\xcc\xaf'.decode()
    }
    bad_chars = {
        # '\ue548',
        # '\ueb81',
        # '\ue727',
        # '\ue40a'
        # 'r'
    }
    ignore = ['sp£ka']
    # debug = '\'spalten\''
    num_boxes = 0
    boxes = set()

    for i, (p1, p2) in tqdm(list(enumerate(zip(iter_block_items(doc1), iter_block_items(doc2))))):
        for char in p2.text:
            if ud.category(char) == 'Co':
                num_boxes += 1
                boxes.add(char)
        if p1.text == p2.text:
            continue
        for p1text, p2text in zip(p1.text.split(" "), p2.text.split(" ")):
            # filters out the changes made
            changes = [(code[0], code[2:]) for code in difflib.ndiff(p1text, p2text)]
            changes_chars = [code[1] for code in changes]

            # search for the pattern of removing a char (likely an unknown unicode char with a valid titus match), and then adding 1 or more chars.
            # which in changes_sequence looks like a minus followed by some number of pluses '-++' or '--++++' or something similar.
            changes_sequence = "".join([code[0] for code in changes])
            pattern = re.compile(r'(-+)(\++)')
            matches = [(match.start(1), match.end(1), match.end(2)) for match in pattern.finditer(changes_sequence)]

            # for every match we have to add to the mapping
            for start, middle, end in matches:
                # identify the broken one (the minus) and the characters replacing it (the pluses after it)
                broken_char = "".join(changes_chars[start:middle])
                fixed_chars = "".join(changes_chars[middle:end])

                if broken_char in ignore:
                    continue

                if broken_char in bad_chars:
                    j_text = "|\t"
                    print(broken_char.encode('unicode_escape').decode(), fixed_chars, sep=" --> text")
                    print(j_text + j_text.join([c.replace("\t", " ") for _, c in changes[start - 10:end + 10]]) + "|", j_text + j_text.join([c for c, _ in changes[start - 10:end + 10]]) + "|", sep="\n")
                    breakpoint()
                broken2fixed[broken_char] = fixed_chars
    breakpoint()
    # {broken.encode(): fixed.encode() for broken, fixed in broken2fixed.items()}
    pass


def find_translate_greek(doc):
    # map the original greek to the unicode
    # segments_font2unicode = {}
    # segments_paragraph_counts = []
    # segments_paragraph_sequences = []

    for p in tqdm(iter_block_items(doc), total=9642):
        # we count up how many times we encounter each translation
        # paragraph_counts = defaultdict(lambda: 0)
        # paragraph_sequences = []
        for run in p.runs:
            # if run.font.name == "TITUS Cyberbit Basic":
            #     breakpoint()
            if run.font.name == "Greek":
                # modify the greek font text to be greek unicode text. using the ▒ character to indicate that we do not have a mapping for that yet.
                new_text = "".join([greek_font2unicode.get(char, "▒") for char in run.text])

                # set the new text and remove the greek font
                run.text = new_text

                # set the translation from font -> unicode
                # segments_font2unicode[run.text] = new_text

                # count how many times we encountered this text to ensure we aren't mangling non-greek text later.
                # paragraph_counts[run.text] += 1
                # paragraph_sequences.append(run.text)
            else:
                # print((run.text, replace_using_regex_titus(run.text, titus_font2unicode)))
                run.text = replace_using_regex_titus(run.text, titus_font2unicode)
                # run.font.name = None
                # for key, value in titus_font2unicode.items():
                #     if key not in run.text:
                #         continue
                #     run.text = run.text.replace(key, value)
                pass
            run.font.name = None
        # for key, value in titus_font2unicode.items():
        #     if key in p.text:
        #         for run in p.runs:
        #             if key in run.text:
        #                 run.text = run.text.replace(key, value)
        # segments_paragraph_counts.append(dict(paragraph_counts))
        # segments_paragraph_sequences.append(paragraph_sequences)
    # return segments_font2unicode, segments_paragraph_counts, segments_paragraph_sequences
    return doc


def translate_replace_greek(doc, segments_font2unicode, segments_paragraph_counts, segments_paragraph_sequences):
    for i, (p, paragraph_counts, paragraph_sequences) in enumerate(zip(doc.paragraphs, segments_paragraph_counts, segments_paragraph_sequences)):
        if len(paragraph_counts) == 0:
            continue
        num_changes = 0
        for run in p.runs:
            for font_text, count in paragraph_counts.items():
                if font_text in run.text:
                    new_text = segments_font2unicode[font_text]
                    run.text = run.text.replace(font_text, new_text, count)
                    num_changes += count
        if num_changes != sum(paragraph_counts.values()):
            breakpoint()
        # breakpoint()
    # breakpoint()
    return doc


def main():
    document = Document('data_nil/NIL (original).docx')
    document_amb = Document('data_nil/NIL (AMB edited, 6-25, CMMC edited 7-07).docx')

    # find_bad_titus(document, document_amb)

    # they have to have the same number of paragraphs otherwise I cannot continue
    # assert len(document_amb.paragraphs) == len(document.paragraphs)

    # segments_font2unicode, segments_paragraph_counts, segments_paragraph_sequences = find_translate_greek(document)
    # new_doc = translate_replace_greek(document_amb, segments_font2unicode, segments_paragraph_counts, segments_paragraph_sequences)

    # breakpoint()
    find_translate_greek(document).save('data_nil/NIL (edited).docx')
    pass


if __name__ == '__main__':
    main()
    pass
